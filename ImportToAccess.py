#!/usr/bin/python

"""Created by Moises Macias. Last Updated: 12/03/2014
Purpose of Program: Import information from an NNS RFQ .txt file located in Q:\Msc NNS Docs\RFQs from Exostar\00.txt
In Case of Error: Contact Duy Nguyen and Command Prompt will notify what line the error has occurred on
"""

from lxml import etree
import pyodbc
from docx import Document

access = pyodbc.connect(r'DRIVER={Microsoft Access Driver (*.mdb)};DBQ=\\fw01sql001\Pentagon\NNS Database.mdb')

accessc = access.cursor()
#sql = pyodbc.connect('DSN=SQL Connection')
sql = pyodbc.connect(driver='{SQL Server}', server='FW01SQL001',  
                      database='LINKTABLES', Trusted_Connection='yes')
sqlc = sql.cursor()

def addPart(tag,text):
    addtoaccess = "update NNS_PARTS set "+tag+"='%s' where Part='%s'" % (text,num)
    accessc.execute(addtoaccess)
    access.commit()    
#def addPentPartExt(text,number,place):
#    addtostockext = "update dbo.STOCKEXT set FIELD%s='%s' where PARTNUMBER='%s'" % (place,text[0:29],number)
#    sqlc.execute(addtostockext)
#    sql.commit()    
#def addPentPart(field,text,part):
#    addtostock = "update dbo.STOCK set "+field+"='%s' where PARTNUMBER='%s'" % (text,part)
#    sqlc.execute(addtostock)
#    sql.commit()

print ("Parsing Commenced...")

adding = 0
c = 0
q = 0
n = 1
m = 0
dev = 0
revision = []
deliver = []


root = etree.Element("RFQPage")
 
filename = open('Q:/Msc NNS Docs/RFQs from Exostar/00.txt', encoding="utf8")
#filename = open('C:/NNS RFQ Temp/00.txt', encoding="utf8")
 
with filename as a:
    listParse = [line.rstrip('\n') for line in a]
     
searchText = ['Quotation Deadline', 'Purchaser ']
BeginParse = ['RFQ ', 'Date:']
rfqPage = ['RFQ Item ', 'Material Number', 'Revision ', 'RFQ Qty', 'Unit', 'Delivery Date', 'Description', 'Noun Name', 'Material', 'Size', 'Type', 'Modifier', 'Remarks', 'Special Conditions','Material Control', 'Material Hardness', 'Operating Pressure', 'National Stock Number', 'Manufacturer', 'Supplemental Description', 'First Tier Specification', 'Specification MOD', 'Latest Approved Revision', 'Classification Data', 'Sheet Number', 'Sheet Suffix', 'Reference Specification', 'Coded Note Description','Specification Doc Type','Specification','Inspect At']
delivcodes = ['12-4','45-14','12-26','76-9M','17-20','60-84','76-9AL','76-151','76-150','37-1','70-20','76-9ER','12-19','76-22B','37-3','9Z','C-85','76-9M','12-19','76-9KZ','37-3','76-20','76-9BG','112-1','76-9LU','S-07','76-9C','45-14','I-24244','C-85','C-22','S-07','12-19','37-19','37-3','1-7','17-20','12-26','76-20','76-9C','45-14','H1010','LI113','LI030','C3040','C2990','C2740','C2560','LI006','C2350','IC190','Q3010','V2010','LI002','C3040','LI117','C2990','H1702','IC033','IC040','M2050','M3010','C3040','LI117','LI002','C2480','D1990','M1350','N2940','ICC14','W4020','Y1920','W2010','C2530','SH1010','C2550','D1190']

FolderDesc =[]

for idx, line in enumerate(listParse):
    for d in BeginParse:
        if line.startswith(d):
            if 'RFQ' in d:
                if 14 <= len(line) <= 17:
                    root.text = (listParse[idx].strip("RFQ  "))
            elif 'Date:' in d:
                etree.SubElement(root, d.strip(":")).text = listParse[idx-1]
            else: 
                pass
    for s in searchText:
        if s in line:
            etree.SubElement(root, line.replace(" ", "").replace(":", "")).text = listParse[idx-1]
    for r in rfqPage:
        if line.startswith(r):
            if 'RFQ Item' in line:
                c = c + 1
            if 'Manufacturer' in line:
                MFR = listParse[idx].replace("Manufacturer's Part No. (PIN/VPN) ","") 
                Cage = MFR + listParse[idx+1] 
                if Cage.find('Cage Code') > 0:
                    etree.SubElement(root, r.replace(" ", "")).text = MFR[0:len(MFR)-6]
                else:
                    etree.SubElement(root, r.replace(" ", "")).text = MFR
            if 'First Tier Specification' in line:
                etree.SubElement(root, r.replace(" ","")).text = listParse[idx-1] + " " + listParse[idx].strip(r)
                continue
            if 'Specification MOD' in line:
                specmod = []
                for l in listParse[idx+1::]:
                    #specmod.append(l)
                    if 'RFQ :' in l:
                        continue;
                    if 'Latest' in l:
                        break;
                    specmod.append(l)
                specmod.pop(-1)
                joinedspecmod = ",".join(specmod)
                etree.SubElement(root, r.replace(" ", "")).text = joinedspecmod
                continue
            if 'Noun Name' in line:
                if 'PACKING' in listParse[idx-1]:
                    for index, i in enumerate(listParse[idx2+1::]):
                        if 'Material Control Level' in i:
                            DescStr = (listParse[idx-1])
                            break
                        if 'Modifier' in i:
                            if (listParse[idx2+1::][index-1]) == 'O-RING':
                                DescStr = (listParse[idx2+1::][index-1])
                            else:
                                DescStr = (listParse[idx-1])
                            break
                else:
                    DescStr =(listParse[idx-1])
                FolderDesc.append(DescStr.replace(" ", "_"))
            if 'Supplemental Description' in line:
                SuppDesc = []
                SuppDesc.append(listParse[idx].replace('Supplemental Description',""))
                idx2 = idx
                for i in enumerate(listParse[idx+1::]):
                    SuppDesc.append(listParse[idx2+1].strip())
                    if 'First Tier Specification Doc Type' in i:
                        SuppDesc.pop(-1)
                        break
                    if 'Specification Doc Type' in i:
                        SuppDesc.pop(-1)
                        break
                    if 'INDUSTRY STD' in i:
                        break
                    if 'Q3010-FIRST ARTICLE TEST-APVL' in i:
                        break
                    if 'SHOCK TEST REPORT-APPROVED' in i:
                        break
                    if 'VENDOR DWG-APPROVED' in i:
                        break
                    if 'NAVSHIP/SEA DWG' in i:
                        break
                    if 'NOT APPLICABLE' in i:
                        break
                    if 'NNS RECEIPT INSPECTION' in i:
                        break;
                    if 'FOUNDATION/INFORMATION DWG-APPROVED' in i:
                        break;
                    if 'LEAD YARD DWG' in i:
                        break;
                    if 'NNS End Use' in i:
                        SuppDesc.pop(-1)
                        break;
                    if 'Inspect By' in i:
                        SuppDesc.pop(-1)
                        break;
                    idx2 = idx2 + 1
                SuppDesc.pop(-1)
                #suppdesc = " ".join(SuppDesc)
                suppdesc = ""
                for x in SuppDesc:
                    if x is None or x == "":
                        suppdesc = suppdesc + x
                    else:
                        suppdesc = suppdesc + "<br>" + x 
                etree.SubElement(root, r.replace(" ","")).text = suppdesc
                continue
            if 'Inspect At' in line:
                Appendices = []
                ADesc = []
                AURL = []
                idx2 = idx
                for i in listParse[idx::]:
                    if i.startswith('http:'):
                        if listParse[idx2-1].startswith('Line Item Descriptions'):
                            ARev = listParse[idx2-3]
                            if ARev.startswith('REV:'):
                                ARev = listParse[idx2-4]
                                ARevsplit = ARev.split(' ')
                                AInst = len(ARev) - len(ARevsplit[-1])
                                ARev = ARev[:AInst] + 'Rev ' + ARev[AInst:]
                        elif listParse[idx2-1].startswith('REV:'):
                            ARev = listParse[idx2-2]
                            ARevsplit = ARev.split(' ')
                            AInst = len(ARev) - len(ARevsplit[-1])
                            ARev = ARev[:AInst] + 'Rev ' + ARev[AInst:]
                        else:
                            ARev = listParse[idx2-1]
                        Appendices.append(ARev)
                        AURL.append(listParse[idx2])
                    if 'Applicable Coded Notes' in i:
                        break;
                    if 'RFQ Item' in i:
                        break;
                    idx2 +=  1
                acoded = "~".join(Appendices)
                codedApp = etree.SubElement(root, "Appendices")
                codedApp.text = acoded
                urlapp = "~".join(AURL)
                aurllink = etree.SubElement(root, "AURL")
                aurllink.text = urlapp
            if 'Coded Note Description' in line:
                CodedNote = []
                Desc = []
                URL = []
                idx2 = idx
                for i in listParse[idx::]:
                    if i.startswith('http:'):
                        if listParse[idx2-1].startswith('Line Item Descriptions'):
                            CodedNote.append(listParse[idx2-3])
                        elif listParse[idx2-1].startswith('** CUI'):
                            CodedNote.append(listParse[idx2-4])
                        else:
                            CodedNote.append(listParse[idx2-1])
                        Desc.append(listParse[idx2+1])
                        URL.append(listParse[idx2])
                    if 'RFQ Item' in i:
                        break;
                    idx2 +=  1
                for d in CodedNote : 
                    if d in delivcodes:
                        deliver.append(d)
                joineddeliver = ",".join(deliver)
                deliverables = etree.SubElement(root, "Deliverables")
                deliverables.text = joineddeliver   
                coded = " ".join(CodedNote)
                codednotes = etree.SubElement(root, "CodedNotes")
                codednotes.text = coded
                desc = ";".join(Desc)
                #descript = etree.SubElement(root, "CodedDesc")
                #descript.text = desc
                url = ",".join(URL)
                urllink = etree.SubElement(root, "URL")
                urllink.text = url
            etree.SubElement(root, r.replace(" ","")).text = listParse[idx-1]


import re
from collections import Counter            
l_sorted = Counter(FolderDesc).most_common()
joinedFolderDesc = ""
for i in range(len(l_sorted)):
    l_joined = ""
    l_split = str(l_sorted[i]).split(", ")
    for x in range(2):
        #alphanumeric = [character for character in l_split[x] if character.isalnum()]
        #alphanumeric = "".join(alphanumeric)
        alphanumeric = re.sub('[^.,-a-zA-Z0-9 \n\.]', '', l_split[x])
        if x == 0:
            l_joined += alphanumeric
        else:
            l_joined += ".(" + alphanumeric +")"
    if i == len(l_sorted) - 1:
        joinedFolderDesc += l_joined
    else:
        joinedFolderDesc += l_joined +", "


for child in root:
    if 'Applicable Coded Notes' in child.text:
        root.remove(child)
    elif 'Date Material' in child.text:
        root.remove(child)
    elif 'Reference Part' in child.text:
        root.remove(child)
    elif 'Quotation' in child.text:
        root.remove(child)
    elif child.text == '':
        root.remove(child)
    elif child.text == '0':
        root.remove(child)
    elif child.text == "COG Code":
        root.remove(child)
    elif child.text == "First Tier Specification":
        root.remove(child)
  
Et = etree.ElementTree(root)
Et.write('Q:/Msc NNS Docs/RFQs from Exostar/RFQtext.xml')
#Et.write('C:/NNS RFQ temp/RFQtext.xml')
filename.close()

#file = open("Q:/Msc NNS Docs/RFQs from Exostar/RFQtext.txt", 'w')
#
#file.write(root.tag+'\n')
#file.write(root.text+'\n')
#
#for child2 in root:
#    file.write(str(child2.tag)+'\n')
#    file.write(child2.text+'\n')
#    
#file.close()    
  
print ("Writing Vendor RFQ...")

rfq = root.text
vendorrfq = Document()
vendorrfq.add_paragraph("Vendor RFQ: " + rfq, style = 'Quote').alingment = 1
     
print ("Importing to Access tables...")
    
rfqfind = "select RFQ from dbo.NNS_RFQ where RFQ like '%s'" % (rfq)
sqlc.execute(rfqfind)
row2 = sqlc.fetchone()
#lastnum = "select top 1 DOC_NO from dbo.QUOTE_HDR ORDER BY DOC_NO DESC"
#sqlc.execute(lastnum)
#findlastnum = sqlc.fetchone()
#rowtostr = str(findlastnum[0])
#docnum = int(rowtostr)+1
#userno = str(docnum).zfill(5)
cleartable = "TRUNCATE TABLE NNS_RFQ_Folder_Label"
sqlc.execute(cleartable)
sql.commit()
rfqadd = "insert into dbo.NNS_RFQ_Folder_Label(Closed,RFQ,Buyer,TotalItems) values ('0',%s,'-',%s)" % (rfq,c)
sqlc.execute(rfqadd)
sql.commit()
adddead = "update dbo.NNS_RFQ_Folder_Label set Descr='%s' where RFQ='%s'" % (joinedFolderDesc,rfq)
sqlc.execute(adddead)
sql.commit()
if row2 == None:
    print ('No RFQ Found. Adding New Entry...')
    rfqadd = "insert into dbo.NNS_RFQ(Closed,RFQ,Buyer,TotalItems) values ('0',%s,'-',%s)" % (rfq,c)
    sqlc.execute(rfqadd)
    sql.commit()
    print ("RFQ: "+rfq+" Successfully Added...")
    #internum = "insert into dbo.QUOTE_HDR(DOC_NO,USER_DOC,DOC_TYPE,DOC_CATEGORY,DOC_STATUS,ACCTNO,SUBC,CCODE,OURCONTACT,VALIDFOR,CUST_REF1,PRIORITY,FOB,FOB_LBL,SMAN1_CODE,CURENCY_BASE,CURENCY_CONV,SIGN_1) values ('%s','%s','Q','QT','4','NOR1','1','0009','DK','30','%s','RG','FOB','F.O.B.','DK','USD','USD','DK')" % (docnum,userno,rfq)
    #sqlc.execute(internum)
    #sql.commit()
    #print ("Pentagon Quote: "+str(docnum)+ " Successfully Added...")
else:
    print ("RFQ: "+rfq+" already Exists...")

    answer = input("Do you want to continue or exit? Enter YES or NO: ")
    realanswer = answer.lower()
    if realanswer == 'yes' or realanswer.startswith('y'):
        print ("Proceeding to update RFQ...")
        pass
    elif realanswer == 'no' or realanswer.startswith('n'):
        print ("Exiting Now...")
        accessc.close()
        access.close()
        exit()
    else:
        print ('Do not recognize response. Exiting...')
        accessc.close()
        access.close()
        exit()
  
findrfqitem = "select count(*) from dbo.NNS_RFQ_ITEMS where RFQ='%s'" % (rfq)
sqlc.execute(findrfqitem)
foundrfqitem = sqlc.fetchone()
#print (foundrfqitem)
if foundrfqitem[0] > 0:
    print(str(foundrfqitem[0]) +" records has been found!")
    deleteanswer = input("You are about to delete item records. Are you sure you want to continue? Enter  YES or NO: ")
    really = deleteanswer.lower()
    if really == "yes" or really.startswith('y'):
        deletesql = "DELETE FROM dbo.NNS_RFQ_ITEMS where RFQ='%s'" % (rfq)
        #deletesql = "DELETE FROM dbo.QUOTE_LINE where DOC_NO='%s'" % ()
        sqlc.execute(deletesql)
        sql.commit()
    else:
        print ("Records were not deleted!...Proceeding to update RFQ...")
        pass
else:
    pass


uc = 0
refjoined = ""
Specstop = 0
firstjoined = ""
Firststop = 0
for t, child in enumerate(root):
    #testp = "Line {}".format(t) + " - " + child.tag
    #print (testp)
    if child.tag == "QuotationDeadline":
        dead = child.text
        vendorrfq.add_paragraph("Quotation Deadline: " + dead)
        rfqfind = "select * from dbo.NNS_RFQ where RFQ like '%s'" % rfq
        sqlc.execute(rfqfind)
        row3 = sqlc.fetchone()
        if row3[4] == "-":
            pass
        else:
            adddead = "update dbo.NNS_RFQ set DateDue='%s' where RFQ='%s'" % (dead,rfq)
            sqlc.execute(adddead)
            sql.commit()
            adddead = "update dbo.NNS_RFQ_Folder_Label set DateDue='%s' where RFQ='%s'" % (dead,rfq)
            sqlc.execute(adddead)
            sql.commit()
            
            #adddeadpent = "update dbo.QUOTE_HDR set ENTER_DATE='%s' where DOC_NO='%s'" % (dead,docnum)
            #sqlc.execute(adddeadpent)
            #sql.commit()
             
    
    elif child.tag == 'PurchaserContact':
        name = child.text
        vendorrfq.add_paragraph("Purchaser: " + name)
        
        findname = "select Name from BUYER_CONTACT where Name like '%s'" % (name)
        accessc.execute(findname)
        namerow = accessc.fetchone()
        newname = "update dbo.NNS_RFQ_Folder_Label set Buyer='%s' where RFQ='%s'" % (name, rfq)
        sqlc.execute(newname)
        sql.commit()
        if namerow == None:
            print ("No Buyer Name found!")

            #Get Initial from name
            xs = (name)
            name_list = xs.split()
            initials = ""
            for name in name_list:  
                initials += name[0].upper()
            
            findname = "select Initials from BUYER_CONTACT where Initials like '%s' Order By Initials DESC" % (initials+"%")
            accessc.execute(findname)
            namerow = accessc.fetchone()

            if namerow == None:
                addBuyer = "insert into BUYER_CONTACT(Initials,Name) values ('%s','%s')" % (initials, child.text)
                accessc.execute(addBuyer)
                access.commit()
            else:
                if namerow[0] == initials:
                    initials = initials+"1"
                    addBuyer = "insert into BUYER_CONTACT(Initials,Name) values ('%s','%s')" % (initials, child.text)
                    accessc.execute(addBuyer)
                    access.commit()
                else:
                    InitCount = int(namerow[0][-1:])+1
                    initials = initials+str(InitCount)
                    addBuyer = "insert into BUYER_CONTACT(Initials,Name) values ('%s','%s')" % (initials, child.text)
                    accessc.execute(addBuyer)
                    access.commit()
                    
            newname = "update dbo.NNS_RFQ set Buyer='%s' where RFQ='%s'" % (initials, rfq)
            sqlc.execute(newname)
            sql.commit()
            print ("New Buyer imported!")
        else:
            print ("Buyer Contact found!")
            findinit = "select Initials from BUYER_CONTACT where Name like '%s'" % (namerow[0])
            accessc.execute(findinit)
            initrow = accessc.fetchone()
            initials = initrow[0]
            
            #print(initrow[0])
            if row3[2] == '-':
                newname = "update dbo.NNS_RFQ set Buyer='%s' where RFQ='%s'" % (initrow[0], rfq)
                sqlc.execute(newname)
                sql.commit()
            else:
                pass
            
    elif child.tag == "PurchaserPhoneNo":
        PurchPhoneNo = child.text
        findbuyer = "select Phone from BUYER_CONTACT where Initials like '%s'" % (initials)
        accessc.execute(findbuyer)
        namerow = accessc.fetchone()
        if namerow[0] != PurchPhoneNo:
            addtoaccess = "update BUYER_CONTACT set Phone ='%s' where Initials='%s'" % (PurchPhoneNo, initials)
            accessc.execute(addtoaccess)
            access.commit()

    elif child.tag == "PurchaserFaxNo":
        PurchFaxNo = child.text
        findbuyer = "select Fax from BUYER_CONTACT where Initials like '%s'" % (initials)
        accessc.execute(findbuyer)
        namerow = accessc.fetchone()
        if namerow[0] != PurchFaxNo:
            addtoaccess = "update BUYER_CONTACT set Fax ='%s' where Initials='%s'" % (PurchFaxNo, initials)
            accessc.execute(addtoaccess)
            access.commit()

    elif child.tag == "PurchaserEmail":
        PurchEmail = child.text
        findbuyer = "select [E-mail] from BUYER_CONTACT where Initials like '%s'" % (initials)
        accessc.execute(findbuyer)
        namerow = accessc.fetchone()
        if namerow[0] != PurchEmail:
            addtoaccess = "update BUYER_CONTACT set [E-mail] ='%s' where Initials='%s'" % (PurchEmail, initials)
            accessc.execute(addtoaccess)
            access.commit()

    elif child.tag == "PurchaserGroup":
        PurchGroup = child.text
        findbuyer = "select [Buyer ID] from BUYER_CONTACT where Initials like '%s'" % (initials)
        accessc.execute(findbuyer)
        namerow = accessc.fetchone()
        if namerow[0] != PurchGroup:
            addtoaccess = "update BUYER_CONTACT set [Buyer ID] ='%s' where Initials='%s'" % (PurchGroup, initials)
            accessc.execute(addtoaccess)
            access.commit()
            
    elif child.tag == "RFQItem":
        rfqitems = child.text
        vendorrfq.add_paragraph("RFQ Item: " + rfqitems).alignment = 1
        findrfqnum = "select count(*) from dbo.NNS_RFQ_ITEMS where RFQ='%s'" % (rfq)
        sqlc.execute(findrfqnum)
        finalrow = sqlc.fetchone()
        #lastquote = "select top 1 LINE from dbo.QUOTE_LINE ORDER BY LINE DESC"
        #sqlc.execute(lastquote)
        #findlastquote = sqlc.fetchone()
        #rowtostr2 = float(findlastquote[0])
        #quotenum = int(rowtostr2)+10
        if finalrow[0] == c:
            pass
        elif finalrow[0] < c+1:
            addrfq = "insert into dbo.NNS_RFQ_ITEMS(RFQ,Item,NNSPart,Qty,Min,EstTargetDate,Awarded) values (%s,%s,'-','0','0','-','0')" % (rfq,rfqitems)
            sqlc.execute(addrfq)
            sql.commit()
            #interquo = "insert into dbo.QUOTE_LINE(DOC_NO,LINE,ACCTNO,SUBC,LINE_TYPE,STATUS,PARTNUMBER,QREQ) values (%s,%s,'NOR1','1','02','O','-','0')" % (docnum,quotenum)
            #sqlc.execute(interquo)
            #sql.commit()
                       
    elif child.tag == 'MaterialNumber':
        Specstop = 0
        num = child.text
        vendorrfq.add_paragraph("Material Number: " + num).alignment = 1
        partfind = "select * from NNS_PARTS where Part LIKE '%s'" % (num)
        accessc.execute(partfind) 
        row = accessc.fetchall()
        #pentpart = "select * from dbo.STOCK where PARTNUMBER LIKE '%"+num+"%'"
        #sqlc.execute(pentpart)
        #pentrow = sqlc.fetchall()
        #print (num)
        print ("NNS Part Updating..."+num)
        if row == []:
            newnum = "insert into NNS_PARTS(Part,Revision,Unit,Description,Noun,Material,Size,Type,Modifier,Remarks,MaterialHardness,OperatingPressure,SpecialConditions,MaterialGroup,RefPart,Spec,Weight,RefSpec,SuppDesc,NSN,CodedNotes,InspectionSpec,Deliverables,Marking,Preservation,QUP,PartSize,BagSize,OPI,IC,ICQ,MarkCode,Die,Label,Seal,Wrap,Bag,Stiffener,Stamp,Cut,Assembly,Pallet,Drawing) values ('%s','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-',No,No,No,No,No,No,No,No,No,No,'-')" % (num)
            accessc.execute(newnum)
            access.commit()
            addpart = "update dbo.NNS_RFQ_ITEMS set NNSPart='%s' where RFQ='%s'and Item='%s'" % (num,rfq,n)
            sqlc.execute(addpart)
            sql.commit()
            #if pentrow == []:
            #    pentnum = "insert into dbo.STOCK(PARTNUMBER,SUBP,DESCRIPTN,ALT1,ALT2,SUPERSEDE_P_F,QTY_INSTOCKUM,QTY_INSTOCKREV,SRCH_NUM,SRCH_PUNCT) values ('%s','1','-','-','-','X','-','-','%s','%s')" % (num,num,num)
            #    sqlc.execute(pentnum)
            #    sql.commit()
            #    extnum = "insert into dbo.STOCKEXT(PARTNUMBER,SUBP,FIELD1,FIELD2,FIELD3,FIELD4,FIELD5,FIELD6,FIELD7,FIELD8,FIELD9,FIELD10,FIELD11,FIELD12,FIELD13,FIELD14,FIELD15,FIELD16,FIELD17,FIELD18,FIELD19,FIELD20) values ('%s',1,'-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-')"  % (num)
            #    sqlc.execute(extnum)
            #    sql.commit()
            #    addpentpart = "update dbo.QUOTE_LINE set PARTNUMBER='%s' where DOC_NO='%s' and LINE='%s'" % (num,docnum,quotenum)
            #    sqlc.execute(addpentpart)
            #    sql.commit()
            #else:
            #    pass
            n += 1
            findpart = "select * from NNS_PARTS where Part LIKE '%s'" % (num)
            accessc.execute(findpart) 
            partrow = accessc.fetchall()
            #findpentp = "select * from dbo.STOCK where PARTNUMBER='%s'" % (num)
            #sqlc.execute(findpentp)
            #pentpartrow = sqlc.fetchone()
            #findpentpart = "select * from dbo.STOCKEXT where PARTNUMBER='%s'" % (num)
            #sqlc.execute(findpentpart)
            #pentextpartrow = sqlc.fetchall()
        elif row:
            addpart = "update dbo.NNS_RFQ_ITEMS set NNSPart='%s' where RFQ='%s' and Item='%s'" % (num,rfq,n)
            sqlc.execute(addpart)
            sql.commit()
            #if pentrow:
            #    addpentpart = "update dbo.QUOTE_LINE set PARTNUMBER='%s' where DOC_NO='%s' and LINE='%s'" % (num,docnum,quotenum)
            #    sqlc.execute(addpentpart)
            #    sql.commit()
            #    findextrow = "select * from dbo.STOCKEXT where PARTNUMBER LIKE '%"+num+"%'"
            #    sqlc.execute(findextrow)
            #    pentextrow = sqlc.fetchall()
            #    if pentextrow == []:
            #         addpentext = "insert into dbo.STOCKEXT(PARTNUMBER,SUBP,FIELD1,FIELD2,FIELD3,FIELD4,FIELD5,FIELD6,FIELD7,FIELD8,FIELD9,FIELD10,FIELD11,FIELD12,FIELD13,FIELD14,FIELD15,FIELD16,FIELD17,FIELD18,FIELD19,FIELD20) values ('%s',1,'-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-')"  % (num)
            #         sqlc.execute(addpentext)
            #         sql.commit()
            #   else:
            #        pass
            #elif pentrow == []:
            #    pentnum = "insert into dbo.STOCK(PARTNUMBER,SUBP,DESCRIPTN,ALT1,ALT2,QTY_INSTOCKUM,QTY_INSTOCKREV,SRCH_NUM,SRCH_PUNCT) values ('%s','1','-','-','-','-','-','%s','%s')" % (num,num,num)
            #    sqlc.execute(pentnum)
            #    sql.commit()
            #    extnum = "insert into dbo.STOCKEXT(PARTNUMBER,SUBP,FIELD1,FIELD2,FIELD3,FIELD4,FIELD5,FIELD6,FIELD7,FIELD8,FIELD9,FIELD10,FIELD11,FIELD12,FIELD13,FIELD14,FIELD15,FIELD16,FIELD17,FIELD18,FIELD19,FIELD20) values ('%s',1,'-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-','-')"  % (num)
            #    sqlc.execute(extnum)
            #    sql.commit()
            #    addpentpart = "update dbo.QUOTE_LINE set PARTNUMBER='%s' where DOC_NO='%s' and LINE='%s'" % (num,docnum,quotenum)
            #    sqlc.execute(addpentpart)
            #    sql.commit()
            #lse: 
            #    pass
            n += 1
            findpart = "select * from NNS_PARTS where Part LIKE '%s'" % (num)
            accessc.execute(findpart) 
            partrow = accessc.fetchall()
            #findpentp = "select * from dbo.STOCK where PARTNUMBER='%s'" % (num)
            #sqlc.execute(findpentp)
            #pentpartrow = sqlc.fetchone()
            #findpentpart = "select * from dbo.STOCKEXT where PARTNUMBER='%s'" % (num)
            #sqlc.execute(findpentpart)
            #pentextpartrow = sqlc.fetchall()
        else:
            findpart = "select * from NNS_PARTS where Part LIKE '%s'" % (num)
            accessc.execute(findpart) 
            partrow = accessc.fetchall()
            #findpentp = "select * from dbo.STOCK where PARTNUMBER='%s'" % (num)
            #sqlc.execute(findpentp)
            #pentpartrow = sqlc.fetchone()
            #findpentpart = "select * from dbo.STOCKEXT where PARTNUMBER='%s'" % (num)
            #sqlc.execute(findpentpart)
            #pentextpartrow = sqlc.fetchall()
                 
        
        
    elif child.tag == "Revision":
        rev = child.text
        revtag = child.tag
        if partrow[0][1] == "-" or partrow[0][1] != rev:
            addPart(revtag,rev)
            #addPentPart("QTY_INSTOCKREV",rev,num)
        #elif pentpartrow[17] == "-" or pentpartrow[17] != rev:
            #addPentPart("QTY_INSTOCKREV",rev,num)
        else:
            pass
                   
    elif child.tag == "Unit":
        unit = child.text
        unittag = child.tag
        vendorrfq.add_paragraph(unit)
        if partrow[0][2] == "-" or partrow[0][2] != unit:
            addPart(unittag,unit)
            #addPentPart("QTY_INSTOCKUM",unit,num)
        #elif pentpartrow[15] == "-" or pentpartrow[15] != unit:
            #addPentPart("QTY_INSTOCKUM",unit,num)
        else:
            pass
                              
    elif child.tag == "RFQQty":
        qty = str(int(float(child.text.replace(',',''))))

        itemnum = vendorrfq.add_paragraph("RFQ Qty: " + qty + " ")
        findnum = "select Item,Qty from dbo.NNS_RFQ_ITEMS where RFQ='%s'" % (rfq)
        sqlc.execute(findnum)
        finalrow2 = sqlc.fetchall()
        addqty = "update dbo.NNS_RFQ_ITEMS set Qty='%s' where RFQ='%s' and Item='%s'" % (qty,rfq,q+1)
        sqlc.execute(addqty)
        sql.commit()
        #addpentqty = "update dbo.QUOTE_LINE set QREQ='%s' where LINE='%s'" % (qty,quotenum)
        #sqlc.execute(addpentqty)
        #sql.commit()
        q += 1 

                   
    elif child.tag == "DeliveryDate":
        deliv = child.text
        vendorrfq.add_paragraph("Delivery Date: " + deliv)
        finddeliv = "select EstTargetDate from dbo.NNS_RFQ_ITEMS where RFQ='%s'" % (rfq)
        sqlc.execute(finddeliv)
        row7 = sqlc.fetchall()
        adddeliv = "update dbo.NNS_RFQ_ITEMS set EstTargetDate='%s' where RFQ='%s' and Item='%s'" % (deliv,rfq,dev+1)
        sqlc.execute(adddeliv)
        sql.commit()
        dev += 1
        
                   
    elif child.tag =="Description":
        desc = child.text
        desctag = child.tag
        vendorrfq.add_paragraph("Description: " + desc)
        if partrow[0][3] == "-" or partrow[0][3] != desc:
            addPart(desctag,desc)
            #addPentPart("DESCRIPTN",desc,num)
        #elif pentpartrow[2] == "-" or pentpartrow[2] != desc:
            #addPentPart("DESCRIPTN",desc,num)
        else:
            pass
                   
    elif child.tag == "NounName":
        noun = child.text
        vendorrfq.add_paragraph("Noun: " + noun)
        if partrow[0][4] == "-" or partrow[0][4] != noun:
            addPart("Noun",noun)
            #addPentPartExt(noun,num,2)
        #elif pentextpartrow[0][3] == "-":
            #addPentPartExt(noun,num,2)
        else:
            pass
               
    elif child.tag == "Material":
        if child.text.isdigit():
            pass
        elif child.text == "STD": 
            pass
        elif "-" in child.text:
            pass
        else:
            mats = child.text
            matstag = child.tag
            vendorrfq.add_paragraph("Material: " + mats)
             
           
    elif child.tag == "Size":
        size = child.text
        sizetag = child.tag
        vendorrfq.add_paragraph("Size: " + size)
        if partrow[0][6] == "-" or partrow[0][6] != size:
            addPart(sizetag,size)
        else:
            pass
                
    elif child.tag == "Type":
        type = child.text
        typetag = child.tag
        vendorrfq.add_paragraph("Type: " + type)
        if partrow[0][7] == "-" or partrow[0][7] != type:
            addPart(typetag,type)
            #addPentPartExt(type,num,4)
        #elif pentextpartrow[0][5] == "-":
            #addPentPartExt(type,num,4)
        else:
            pass
               
    elif child.tag == "Modifier":
        modi = child.text
        moditag = child.tag
        vendorrfq.add_paragraph("Modifier: " + modi)
        if partrow[0][8] == "-" or partrow[0][8] != modi:
            addPart(moditag,modi)
            #addPentPartExt(modi,num,5)
        #elif pentextpartrow[0][6] == "-":
            #addPentPartExt(modi,num,5)
        else:
            pass
                   
    elif child.tag == "Remarks":
        rem = child.text
        remtag = child.tag
        vendorrfq.add_paragraph("Remarks: " + rem)
        if partrow[0][9] == "-" or partrow[0][9] != rem:
            addPart(remtag,rem)
            #addPentPartExt(rem,num,6)
        #elif pentextpartrow[0][7] == "-":
            #addPentPartExt(rem,num,6)
        else:
            pass
                 
    elif child.tag == "MaterialHardness":
        mh = child.text
        mhtag = child.tag
        vendorrfq.add_paragraph("Material Hardness: "+ mh)
        if partrow[0][10]  == "-" or partrow[0][10] != mh:
            addPart(mhtag,mh)
            #addPentPartExt(mh,num,7)
        #elif pentextpartrow[0][8] == "-":
            #addPentPartExt(mh,num,7)
        else: 
            pass
                 
    elif child.tag == "SpecialConditions":
        sc = child.text
        sctag = child.tag
        vendorrfq.add_paragraph("Special Conditions: " + sc)
        if partrow[0][11] == "-" or partrow[0][11] != sc:
            addPart(sctag,sc)
            #addPentPartExt(sc,num,8)
        #elif pentextpartrow[0][9] == "-":
            #addPentPartExt(sc,num,8)
        else:
            pass    
                 
    elif child.tag == "OperatingPressure":
        op = child.text
        optag = child.tag
        vendorrfq.add_paragraph("Operating Pressure: " + op)
        if partrow[0][12] == "-" or partrow[0][12] != op:
            addPart(optag,op)
            #addPentPartExt(op,num,9)
        #elif pentextpartrow[0][10] == "-":
            #addPentPartExt(op,num,9)
        else:
            pass
                              
    elif child.tag == "MaterialControl":
        mat = child.text
        if partrow[0][13] == "-" or partrow[0][13] != mat:
            addPart("MaterialGroup",mat)
            #addPentPartExt(mat,num,10)          
        #elif pentextpartrow[0][11] == "-":
            #addPentPartExt(mat,num,10)
        else:
            pass
                        
    elif child.tag == "Manufacturer":
        ref = child.text
        vendorrfq.add_paragraph("Manufacturer's Part Number (PIN/VPN): " + ref)
        if partrow[0][14] == "-" or partrow[0][14] != ref:
            addPart("RefPart",ref)
            #addPentPartExt(ref,num,11)
        #elif pentextpartrow[0][12] == "-":
            #addPentPartExt(ref,num,11)
        else:
            pass
                   
    elif child.tag == "FirstTierSpecification":
        first = []
        firstjoined = ""
        SheetNumber =""
        SheetSuffix =""
        SpecificationMOD=""
        fj = 0
        specrev = ""
        for f in root[t::]:                
                #print ("{}".format(fj) + " - " + f.tag)
                if f.tag == "SpecificationDocType" or f.tag == "CodedNotes" or f.tag == "Deliverables" or f.tag == "ReferenceSpecification" or f.tag == "Appendices" or f.tag == "InspectAt":
                    if SheetNumber !="":
                        firstjoined = " " + firstjoined + "Sheet Number " + SheetNumber + " "
                    if SheetSuffix !="":
                        firstjoined = " " + firstjoined + "Sheet Suffix " + SheetSuffix + " "
                    firstjoined = firstjoined + "Rev " + specrev
                    Firststop = 1
                    break
                if f.text.find("Doc Ty") > 0:
                    firstjoined = "-"
                    continue
                if fj == 0:
                    firstjoined = f.text + " "
                    fj =+ 1
                    continue
                if "SpecificationMOD" in f.tag:
                    SpecificationMOD = f.text
                    continue
                if "Specification" in f.tag:
                    continue
                if "LatestApprovedRevision" in f.tag:
                    if SpecificationMOD =="":
                        specrev = f.text
                    else:
                        specrev = f.text + ", MOD = " + SpecificationMOD
                    continue
                if "SheetNumber" in f.tag:
                    SheetNumber = f.text
                    continue
                if "SheetSuffix" in f.tag:
                    SheetSuffix = f.text
                    continue

                firstjoined = firstjoined + f.text + " "
        rfsupp = ",".join(first)
        #joinedfirst = rfsupp.replace("'", "")
        if (len(first)) == 5:
            joinedfirst =(first[2]+","+first[4]+", REV "+first[3])
        if (len(first)) == 4:
            joinedfirst =(first[1]+","+first[3]+", REV "+first[2])
        if (len(first)) == 3:
            joinedfirst =(first[0]+","+first[2]+", REV "+first[1])
        if (len(first)) == 2:
            joinedfirst = (first[0]+" REV "+first[1])    
        vendorrfq.add_paragraph("First Tier Specification: " + firstjoined)
        if partrow[0][15] == "-" or partrow[0][15] != firstjoined:
            addPart("Spec",firstjoined)
            #addPentPart("ALT2",joinedfirst[0:29],num)
        #elif pentpartrow[4] == "-" or pentpartrow[4] != joinedfirst:
            #addPentPart("ALT2",joinedfirst[0:29],num)
        else:
            pass
                   
    elif child.tag == "ReferenceSpecification" or child.tag == "SpecificationDocType":
        ref = []
        refjoined =""
        SheetNumber =""
        SheetSuffix =""
        rj = 0
        specrev = ""
        for l in root[t::]:
            
            if Specstop == 0:
                if "DocType" in l.tag:
                    if refjoined != "":
                        if SheetNumber !="":
                            refjoined = refjoined + "Sheet Number " + SheetNumber + " "
                        if SheetSuffix !="":
                            refjoined = refjoined + "Sheet Suffix " + SheetSuffix + " "
                        refjoined = refjoined + "Rev " + specrev + " | "
                    DocType = l.text    
                    continue
                if l.text == DocType:
                    continue
                if "Specification" in l.tag:
                    pass
                if "LatestApprovedRevision" in l.tag:
                    specrev = l.text
                    continue
                if "SheetNumber" in l.tag:
                    SheetNumber = l.text
                    continue
                if "SheetSuffix" in l.tag:
                    SheetSuffix = l.text
                    continue
                if l.tag == "CodedNotes" or l.tag == "Deliverables" or l.tag == "Appendices":
                    refjoined = refjoined + "Rev " + specrev
                    Specstop = 1
                    break
            if Specstop == 1:
                break
            refjoined = refjoined + l.text + " "
                
        joinedref = ",".join(ref)
        if refjoined !="":
            vendorrfq.add_paragraph("Reference Specification: " + refjoined)
        if partrow[0][16] == "-" or partrow[0][16] != refjoined:
            if len(refjoined) != 0:
                addPart("RefSpec",refjoined)
            #addPentPartExt(joinedref[0:29],num,12)
        #elif pentextpartrow[0][13] == "-":
            #addPentPartExt(joinedref[0:29],num,12)
        else:
            pass
        
               
    elif child.tag == "SupplementalDescription":
        supp = child.text
        finalsupp = ""
        count = 0
        for k in supp:
            if count < 63000:
                finalsupp += k
                count += 1
            else:
                pass
        vendorrfq.add_paragraph("Supplemental Description: " + finalsupp)
        if partrow[0][18] == "-" or partrow[0][18] != finalsupp:
            rfsupp = finalsupp.replace("'", "")
            addPart("SuppDesc",rfsupp)
            #addPentPartExt(finalsupp[0:29],num,1)
        #elif pentextpartrow[0][2] == "-":
            #addPentPartExt(finalsupp[0:29],num,1)
        else:
            pass
                  
    elif child.tag == "NationalStockNumber":
        natnum = child.text
        vendorrfq.add_paragraph("National Stock Number: " + natnum)
        if partrow[0][19] == "-" or partrow[0][19] != natnum:
            addPart("NSN",natnum)
            #addPentPart("ALT1",natnum,num)
        #elif pentpartrow[3] == "-" or pentpartrow[3] != natnum:
            #addPentPart("ALT1",natnum,num)
        else:
            pass
                        
    elif child.tag == "CodedNotes":
        code = child.text
        codetag = child.tag
        eachcode = code.split()
        for codename in eachcode:
            getcode = "select * from Code_Translations where Code LIKE '%s'" % codename
            accessc.execute(getcode)
            row4 = accessc.fetchone()
            if row4 == None:
                newcode = "insert into Code_Translations(Code) values ('%s')" % codename
                accessc.execute(newcode)
                access.commit()
            else:
                pass
        if partrow[0][20] == "-" or partrow[0][20] != code:
            addPart(codetag,code)
            #addPentPartExt(code[0:29],num,13)
        #elif pentextpartrow[0][14] == "-":
            #addPentPartExt(code[0:29],num,13)
        else:
            pass
                                       
    elif child.tag == "URL":
        u = 0
        urll = child.text.split(',')
        for oldcode in eachcode:
            geturl = "select URL from Code_Translations where Code LIKE '%s'" % (oldcode)
            accessc.execute(geturl)
            row5 = accessc.fetchone()
            if row5[0] == None or row5[0] != urll[u]:
                newurl = "update Code_Translations set URL='%s' where Code='%s'" % (urll[u],oldcode)
                accessc.execute(newurl)
                access.commit() 
            else:
                pass
            u += 1
                           
    elif child.tag == "CodedDesc":
        d = 0
        codedesc = child.text.split(';')
        for veryoldcode in eachcode:
            vendorrfq.add_paragraph("Coded Notes: " + veryoldcode + " - " + codedesc[d])
            getdesc = "select Description from Code_Translations where Code LIKE '%s'" % veryoldcode
            accessc.execute(getdesc)
            row6 = accessc.fetchone()
            if row6[0] == None or row6[0] != codedesc[d]:
                newcodedesc = "update Code_Translations set Description='%s' where Code='%s'" % (codedesc[d],veryoldcode)
                accessc.execute(newcodedesc)
                access.commit()
            else:
                pass
            d += 1 
        vendorrfq.add_page_break()
                          
    elif child.tag == "Deliverables":
        deliverable = child.text
        deltag = child.tag
        if partrow[0][22] == "-" or partrow[0][22] != deliverable:
            addPart(deltag,deliverable)
            #addPentPartExt(deliverable[0:29],num,19)
        #elif pentextpartrow[0][20] == "-":
            #addPentPartExt(deliverable[0:29],num,19)
        else:
            pass

                      
    elif child.tag == "Date":
        date = child.text
        datercvd = "update dbo.NNS_RFQ set DateRcvd='%s' where RFQ='%s'" % (date,rfq)
        sqlc.execute(datercvd)
        sql.commit()
        datercvd = "update dbo.NNS_RFQ_Folder_Label set DateRcvd='%s' where RFQ='%s'" % (date,rfq)
        sqlc.execute(datercvd)
        sql.commit() 
        #dateenter = "update dbo.QUOTE_HDR set ENTER_DATE='%s', DOC_DATE='%s' where DOC_NO='%s'" % (date,date,docnum)
        #sqlc.execute(dateenter)
        #sql.commit()

    elif child.tag == "Appendices":
        appen = child.text
        appentag = child.tag
        eachappen = appen.split('~')
        for appenname in eachappen:
            getappen = "select * from Appendices where Appendice LIKE '%s'" % appenname
            accessc.execute(getappen)
            row4 = accessc.fetchone()
            if row4 == None:
                newappen = "insert into Appendices(Appendice) values ('%s')" % appenname
                accessc.execute(newappen)
                access.commit()
            else:
                pass
            nappen = n
            addappen = "update dbo.NNS_RFQ_ITEMS set Appendices='%s' where RFQ='%s'and Item='%s'" % (appen,rfq,nappen-1)
            sqlc.execute(addappen)
            sql.commit()

    elif child.tag == "AURL":
        au = 0
        aurll = child.text.split('~')
        for oldappen in eachappen:
            getaurl = "select URL from Appendices where Appendice LIKE '%s'" % (oldappen)
            accessc.execute(getaurl)
            rowaurl = accessc.fetchone()
            if rowaurl[0] == None or rowaurl[0] != aurll[au]:
                anewurl = "update Appendices set URL='%s' where Appendice='%s'" % (aurll[au],oldappen)
                accessc.execute(anewurl)
                access.commit() 
            else:
                pass
            au += 1
    else:
        pass


       
accessc.close()
access.close()
sqlc.close()
sql.close()
#vendorrfq.save("Q:/Msc NNS Docs/RFQs from Exostar/"+rfq+".docx")

from subprocess import Popen
Popen("Q:/LABELS/Programs/NNS_RFQ_Folder_Label.bat")

print ("Import Complete")
#import ctypes  # An included library with Python install.
#ctypes.windll.user32.MessageBoxW(0, "", "Import Completed", 1)

#print ("Press any key to continue...")
